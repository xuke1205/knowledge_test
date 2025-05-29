"""
Document Processing Module for Knowledge Base Construction.

This module handles the loading, cleaning, and preprocessing of different document formats.
"""

import os
import re
import logging
from typing import List, Dict, Any, Optional, Tuple

# Import document loaders
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from unidecode import unidecode

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document loading, cleaning, and preprocessing for knowledge base construction."""
    
    def __init__(self, noise_patterns: Optional[List[str]] = None):
        """
        Initialize the document processor.
        
        Args:
            noise_patterns: List of regex patterns to filter out noise (e.g., headers, footers)
        """
        self.noise_patterns = noise_patterns or [
            r'Page \d+ of \d+',        # Common page numbers
            r'\d{1,2}/\d{1,2}/\d{2,4}', # Common date formats
            r'^Document ID:.*$',        # Document IDs
            r'Confidential',            # Confidentiality notices
            r'©\s?\d{4}.*',             # Copyright notices
            r'^\s*\d+\s*$',             # Standalone numbers (like page numbers)
            r'^\s*-\s*\d+\s*-\s*$',     # Another page number format
        ]
        logger.info("Document processor initialized")
    
    def load_document(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load document content from different file formats.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            A tuple containing:
                - The extracted text content
                - Metadata about the document
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # Extract content based on file type
        if ext == '.pdf':
            return self._load_pdf(file_path)
        elif ext == '.docx':
            return self._load_docx(file_path)
        elif ext == '.txt':
            return self._load_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _load_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from PDF files."""
        logger.info(f"Loading PDF: {file_path}")
        
        try:
            doc = fitz.open(file_path)
            
            # Extract metadata
            metadata = {
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'keywords': doc.metadata.get('keywords', ''),
                'page_count': len(doc),
                'file_size': os.path.getsize(file_path),
                'file_name': os.path.basename(file_path),
                'file_type': 'pdf'
            }
            
            # Extract text content
            text_content = []
            for page_num, page in enumerate(doc):
                text = page.get_text()
                text_content.append(f"Page {page_num + 1}:\n{text}")
            
            full_text = "\n".join(text_content)
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error loading PDF file {file_path}: {e}")
            raise
    
    def _load_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from DOCX files."""
        logger.info(f"Loading DOCX: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract metadata
            core_properties = doc.core_properties
            metadata = {
                'title': core_properties.title or '',
                'author': core_properties.author or '',
                'subject': core_properties.subject or '',
                'keywords': core_properties.keywords or '',
                'file_size': os.path.getsize(file_path),
                'file_name': os.path.basename(file_path),
                'file_type': 'docx'
            }
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs]
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    table_text.append(" | ".join(row_text))
                tables_text.append("\n".join(table_text))
            
            # Combine all text
            full_text = "\n\n".join(paragraphs + tables_text)
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error loading DOCX file {file_path}: {e}")
            raise
    
    def _load_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Load plain text files."""
        logger.info(f"Loading TXT: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Basic metadata for text files
            metadata = {
                'file_size': os.path.getsize(file_path),
                'file_name': os.path.basename(file_path),
                'file_type': 'txt'
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            # Try different encodings if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                return text, metadata
            except Exception as e:
                logger.error(f"Error loading TXT file {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"Error loading TXT file {file_path}: {e}")
            raise
    
    def clean_text(self, text: str) -> str:
        """
        Clean the extracted text by removing noise and standardizing format.
        
        Args:
            text: The text to clean
            
        Returns:
            Cleaned text
        """
        logger.info("Cleaning text content")
        
        # Remove noise using defined patterns
        cleaned_text = text
        for pattern in self.noise_patterns:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.MULTILINE)
        
        # Standardize whitespace
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        # Keep original text encoding for non-ASCII characters
        cleaned_text = cleaned_text
        
        # Remove excessive line breaks
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
        
        return cleaned_text.strip()

    def process_document(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Process a document - load and clean its content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            A tuple containing:
                - The cleaned text content
                - Metadata about the document
        """
        # Load the document
        text, metadata = self.load_document(file_path)
        
        # Clean the text
        cleaned_text = self.clean_text(text)
        
        # Add processing metadata
        metadata['original_length'] = len(text)
        metadata['cleaned_length'] = len(cleaned_text)
        metadata['reduction_ratio'] = 1 - (len(cleaned_text) / len(text)) if len(text) > 0 else 0
        
        logger.info(f"Document processed: {os.path.basename(file_path)}")
        
        return cleaned_text, metadata

# Usage example
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Example usage - process different document types
    sample_files = [
        "sample.pdf",
        "sample.docx",
        "sample.txt"
    ]
    
    for file_path in sample_files:
        if os.path.exists(file_path):
            try:
                text, meta = processor.process_document(file_path)
                print(f"Processed {file_path}: {len(text)} chars, {meta.get('file_type')}")
                print(f"First 200 chars: {text[:200]}...")
                print("-" * 50)
            except Exception as e:
                print(f"Error processing {file_path}: {e}")
